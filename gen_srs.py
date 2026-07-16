from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Margins ──
for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# ── Color Palette (Navy and Coral Accent) ──
COLOR_NAVY = RGBColor(27, 67, 50)      # #1B4332 (Primary brand color)
COLOR_CORAL = RGBColor(244, 140, 6)    # #F48C06 (Accent color)
COLOR_TEXT = RGBColor(38, 38, 38)      # #262626 (Charcoal text)
COLOR_MUTED = RGBColor(120, 120, 120)  # Muted grey

# ── Document Defaults (Times New Roman, LTR) ──
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
font.color.rgb = COLOR_TEXT
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

def set_cell_background(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def spacer():
    doc.add_paragraph('')

# ── Stylized Document Builders ──
def h(text, level=1):
    hd = doc.add_heading(text, level=level)
    hd.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in hd.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = COLOR_NAVY
        run.bold = True
        if level == 1:
            run.font.size = Pt(18)
        else:
            run.font.size = Pt(14)
            run.font.color.rgb = COLOR_CORAL
    return hd

def p(text, bold=False, italic=False, color=COLOR_TEXT, size=12):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    return para

def bullet(text):
    para = doc.add_paragraph(style='List Bullet')
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_TEXT
    return para

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Stylized Header
    for i, h_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_background(cell, '1B4332') # Navy header
        set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
        cell.text = ''
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(h_text)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255) # White text
        
    # Styled Rows
    for r_idx, row_data in enumerate(rows):
        row = table.add_row()
        bg_color = 'F8F9FA' if r_idx % 2 == 0 else 'FFFFFF'
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            cell.text = ''
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(val)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.color.rgb = COLOR_TEXT
            
    doc.add_paragraph('')
    return table

# ══════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════
for _ in range(4):
    doc.add_paragraph('')

# Accent line
line_para = doc.add_paragraph()
line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
line_run = line_para.add_run('❖ ❖ ❖')
line_run.font.color.rgb = COLOR_CORAL
line_run.font.size = Pt(16)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Software Requirements Specification\n& Project Quotation')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(26)
run.font.color.rgb = COLOR_NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('E-Commerce Platform with Cash on Delivery (COD) for the Algerian Market')
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
run.font.color.rgb = COLOR_MUTED

for _ in range(2):
    doc.add_paragraph('')

# Main info card
info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info_p.add_run(
    'Developer: Amraoui Mohamed\n'
    'Email: amraouimohamedtaher@gmail.com\n'
    'Client: Algerian E-Commerce Seller\n'
    'Date: June 26, 2026\n'
    'Exchange Rate: 1 USD = 250 DZD\n'
)
run.font.name = 'Times New Roman'
run.font.size = Pt(11)
run.font.color.rgb = COLOR_TEXT

price_p = doc.add_paragraph()
price_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = price_p.add_run('Total Project Price: 60,000 DZD')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(16)
run.font.color.rgb = COLOR_CORAL

doc.add_page_break()

# ══════════════════════════════════════
# CONTENTS
# ══════════════════════════════════════

h('1. General Project Details')
add_table(
    ['Parameter', 'Details'],
    [
        ['Software Name', 'E-Commerce Platform - Cash on Delivery (COD)'],
        ['Lead Developer', 'Amraoui Mohamed'],
        ['Developer Email', 'amraouimohamedtaher@gmail.com'],
        ['Client Name', 'Algerian E-Commerce Seller (Undisclosed)'],
        ['Target Market & Currency', 'Algeria (Wilaya/Commune support / DZD)'],
        ['Agreed Price', '60,000 DZD'],
        ['Exchange Rate', '1 USD = 250 DZD'],
    ]
)

h('2. About the Project')
p('This project is a dedicated, fully customized e-commerce system built for Algerian online sellers to simplify selling products using Cash on Delivery (COD). The system removes the need for mandatory user account registration or complex online payment gateways.')
p('Customers select their items, choose sizes and colors, fill in shipping details (Wilaya, Commune, Address, Phone), and submit their orders immediately. The seller manages these orders via a custom backend dashboard to confirm details with the client over the phone.')

h('3. Main System Features')
p('A. Frontend Customer Interface:', bold=True, color=COLOR_NAVY)
bullet('Fully responsive landing page featuring product displays, steps, store benefits, and a contact section.')
bullet('Catalog page with advanced filtering options (by category, price range, stock availability, and sorting).')
bullet('Product detail page with an interactive gallery, size selector, and color selection.')
bullet('Guest checkout form to complete orders quickly in a few seconds.')
bullet('Dynamic dropdown selection for Algerian Wilayas and Communes to reduce delivery address errors.')

p('B. Dashboard & Inventory Management:', bold=True, color=COLOR_NAVY)
bullet('Centralized dashboard sidebar layout tracking recent orders, stock levels, and sales reports.')
bullet('Complete stock CRUD operations including product main images, gallery uploads, sizes, and colors.')
bullet('Order tracking system with status flows (Pending, Confirmed, Shipped, Completed, Cancelled).')
bullet('Automatic inventory stock restoration when an order status is marked as Cancelled.')
bullet('Real-time notifications using Server-Sent Events (SSE) when a new order is received.')

h('4. Features Breakdown & Pricing Table')
add_table(
    ['#', 'Feature / Component Description', 'Price (DZD)'],
    [
        ['1', 'Homepage UI & Layout design (Hero, Features, Contact)', '8,000 DZD'],
        ['2', 'Product list / catalog page with filters and pagination', '7,000 DZD'],
        ['3', 'Product details page with gallery and variant selector', '8,000 DZD'],
        ['4', 'Guest Checkout flow with dynamic Wilaya/Commune loader', '10,000 DZD'],
        ['5', 'Dashboard core scaffolding (Sidebar, Notifications, Profile)', '7,000 DZD'],
        ['6', 'Stock CRUD management with multi-card product forms', '8,000 DZD'],
        ['7', 'Order processing system (status changes, printable invoices, auto-restoration)', '6,000 DZD'],
        ['8', 'Real-time Server-Sent Events (SSE) notification module', '3,000 DZD'],
        ['9', 'Global theme color customization engine (Context Processor)', '2,000 DZD'],
        ['10', 'Docker environment configuration & developer handover files', '1,000 DZD'],
        ['', 'Total Project Price', '60,000 DZD'],
    ]
)

h('5. Product Create & Order Flow')
p('A. Product Creation Workflow (Seller):', bold=True, color=COLOR_CORAL)
p('The seller navigates to Dashboard → Stock → Create Product. The form is organized into three clean cards: Product Images (main and gallery), Basic Information (name, category, price, description), and Inventory & Variants (quantity, sizes, colors). The form is submitted asynchronously via AJAX, triggering a success popup using SweetAlert.')

p('B. Ordering Workflow (Customer):', bold=True, color=COLOR_CORAL)
p('The customer browses the catalog, views product details, selects their variations (size, color, quantity), fills in their details (name, phone, wilaya, commune, address), and clicks "Order Now". The order is booked instantly, reserving the items from the stock, and pushing a real-time event to the seller\'s dashboard.')

h('6. Included vs. Not Included Services')
add_table(
    ['Included in Project Price ✅', 'Not Included in Project Price ❌'],
    [
        ['Full source code (clean Python/Django/JS)', 'Cloud VPS server hosting fees'],
        ['Database setup (PostgreSQL with migration files)', 'Domain name registration (.com, .dz, etc.)'],
        ['Docker environment containerization files', 'SSL Security Certificate setup fees'],
        ['Full developer documentation (SRS, README, AGENTS)', 'Store branding design or logo design'],
        ['1-month free bug-fix and maintenance support', 'Server deployment and domain binding setup'],
        ['1 remote training session for dashboard usage', 'Digital marketing and advertisements'],
    ]
)

h('7. Server Deployment & Setup (Optional)')
p('This deployment package is optional and is billed separately from the core development project. The client is responsible for purchasing the hosting VPS and domain name prior to starting the setup.')
add_table(
    ['Optional Deployment Services', 'Price (DZD)', 'Notes'],
    [
        ['VPS Setup, Docker installation, & Nginx security configuration', '5,000 DZD', 'Requires server credentials'],
        ['Domain binding & SSL installation', '2,000 DZD', 'Requires domain registrar access'],
        ['Full build deployment, migration, & live testing', '3,000 DZD', 'Final validation'],
        ['Total Setup Price (Optional)', '10,000 DZD', 'Billed separately'],
    ]
)

h('8. Estimated Development Timeline')
add_table(
    ['Development Milestone', 'Estimated Duration'],
    [
        ['Environment setup and core database model setup', '2 Days'],
        ['Frontend pages design and guest checkout system', '5 Days'],
        ['Dashboard management panel and reports interface', '5 Days'],
        ['Order processing, inventory triggers, and SSE events', '3 Days'],
        ['Integration testing, final review, and handover documentation', '2 Days'],
        ['Total Time to Delivery', '17 Working Days'],
    ]
)

spacer()
spacer()

footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('— Prepared by —\nAmraoui Mohamed\namraouimohamedtaher@gmail.com')
run.font.name = 'Times New Roman'
run.font.size = Pt(11)
run.font.color.rgb = COLOR_MUTED

# Save
path = os.path.join(os.path.dirname(__file__), 'SRS_Final.docx')
doc.save(path)
print(f'Saved to {path}')
